from __future__ import annotations

import hashlib
import os
import re
import shutil
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from dotenv import load_dotenv
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_google_genai import (
    ChatGoogleGenerativeAI,
    GoogleGenerativeAIEmbeddings,
)
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader


load_dotenv()

# Change this value whenever indexing logic changes.
# A new collection name prevents old chunks from being reused.
INDEX_VERSION = "v3_no_overlap_no_repeated_paragraphs"
COLLECTION_NAME = f"ragify_{INDEX_VERSION}"


@dataclass
class IndexResult:
    total_files: int
    total_pages: int
    total_chunks: int
    skipped_files: list[str] = field(default_factory=list)
    errors: list[str] = field(default_factory=list)


@dataclass
class AnswerResult:
    answer: str
    sources: list[Document]


class VectorManager:
    def __init__(
        self,
        persist_directory: str = "chroma_db",
    ) -> None:
        api_key = os.getenv("GOOGLE_API_KEY")

        if not api_key:
            raise ValueError(
                "GOOGLE_API_KEY is missing from the .env file."
            )

        self.persist_directory = persist_directory

        self.embeddings = GoogleGenerativeAIEmbeddings(
            model="models/gemini-embedding-001",
            google_api_key=api_key,
        )

        self.vectorstore = self._build_store()

    def _build_store(self) -> Chroma:
        return Chroma(
            collection_name=COLLECTION_NAME,
            persist_directory=self.persist_directory,
            embedding_function=self.embeddings,
        )

    def add_documents(
        self,
        documents: list[Document],
    ) -> None:
        if not documents:
            return

        ids: list[str] = []

        for document in documents:
            source = str(document.metadata.get("source", ""))
            page = str(document.metadata.get("page", ""))
            chunk_index = str(document.metadata.get("chunk_index", ""))

            content_hash = hashlib.sha256(
                document.page_content.encode(
                    "utf-8",
                    errors="ignore",
                )
            ).hexdigest()[:20]

            ids.append(f"{source}-{page}-{chunk_index}-{content_hash}")

        self.vectorstore.add_documents(
            documents=documents,
            ids=ids,
        )

    def similarity_search(
        self,
        query: str,
        k: int = 6,
    ) -> list[Document]:
        return self.vectorstore.similarity_search(
            query,
            k=k,
        )

    def get_all_chunks(self) -> list[Document]:
        result = self.vectorstore.get(
            include=["documents", "metadatas"]
        )

        texts = result.get("documents") or []
        metadatas = result.get("metadatas") or []

        documents: list[Document] = []

        for text, metadata in zip(texts, metadatas):
            documents.append(
                Document(
                    page_content=text or "",
                    metadata=metadata or {},
                )
            )

        documents.sort(
            key=lambda item: (
                str(item.metadata.get("source", "")),
                int(item.metadata.get("page", 0) or 0),
                int(item.metadata.get("chunk_index", 0) or 0),
            )
        )

        return documents

    def clear(self) -> None:
        try:
            self.vectorstore.delete_collection()
        except Exception:
            pass

        self.vectorstore = self._build_store()

    def clear_everything(self) -> None:
        try:
            self.vectorstore.delete_collection()
        except Exception:
            pass

        directory = Path(self.persist_directory)

        if directory.exists():
            try:
                shutil.rmtree(directory)
            except PermissionError:
                # Chroma can keep a file handle open on Windows.
                # Rebuild only the current versioned collection.
                pass

        self.vectorstore = self._build_store()


class RAGPipeline:
    def __init__(self) -> None:
        api_key = os.getenv("GOOGLE_API_KEY")

        if not api_key:
            raise ValueError(
                "GOOGLE_API_KEY is missing from the .env file."
            )

        self.vector_manager = VectorManager()

        self.llm = ChatGoogleGenerativeAI(
            model="gemini-flash-latest",
            google_api_key=api_key,
            temperature=0.1,
        )

        # Zero overlap means neighboring chunks will not repeat text.
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=0,
            separators=[
                "\n\n",
                "\n",
                ". ",
                "? ",
                "! ",
                " ",
                "",
            ],
            length_function=len,
        )

    @staticmethod
    def _clean_text(
        text: str,
    ) -> str:
        if not text:
            return ""

        text = text.replace("\u00ad", "")
        text = re.sub(r"-\s*\n\s*", "", text)
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n[ \t]+", "\n", text)
        text = re.sub(r"\n{3,}", "\n\n", text)

        return text.strip()

    @staticmethod
    def _normalize_for_comparison(
        text: str,
    ) -> str:
        text = text.lower()
        text = re.sub(r"\s+", " ", text)
        text = re.sub(r"[^\w\s]", "", text)

        return text.strip()

    def _remove_repeated_paragraphs(
        self,
        text: str,
        seen_paragraphs: set[str],
    ) -> str:
        """
        Remove paragraphs already seen on an earlier page of the same file.

        This is different from chunk overlap. It handles PDFs that repeat
        the same executive summary or report paragraph on several pages.
        """
        cleaned = self._clean_text(text)

        if not cleaned:
            return ""

        paragraphs = re.split(
            r"\n\s*\n|(?<=[.!?])\s+(?=[A-Z])",
            cleaned,
        )

        kept: list[str] = []

        for paragraph in paragraphs:
            paragraph = paragraph.strip()

            if not paragraph:
                continue

            normalized = self._normalize_for_comparison(paragraph)

            # Very short headings may legitimately repeat.
            if len(normalized) < 35:
                kept.append(paragraph)
                continue

            if normalized in seen_paragraphs:
                continue

            seen_paragraphs.add(normalized)
            kept.append(paragraph)

        return "\n\n".join(kept).strip()

    def _read_pdf(
        self,
        uploaded_file: Any,
    ) -> tuple[list[Document], int]:
        uploaded_file.seek(0)
        reader = PdfReader(uploaded_file)

        documents: list[Document] = []
        readable_pages = 0
        seen_paragraphs: set[str] = set()

        for page_index, page in enumerate(
            reader.pages,
            start=1,
        ):
            raw_text = page.extract_text() or ""

            page_text = self._remove_repeated_paragraphs(
                raw_text,
                seen_paragraphs,
            )

            if not page_text:
                continue

            readable_pages += 1

            documents.append(
                Document(
                    page_content=page_text,
                    metadata={
                        "source": uploaded_file.name,
                        "page": page_index,
                        "location": f"Page {page_index}",
                    },
                )
            )

        uploaded_file.seek(0)

        return documents, readable_pages

    def _read_docx(
        self,
        uploaded_file: Any,
    ) -> tuple[list[Document], int]:
        try:
            from docx import Document as WordDocument
        except ImportError as error:
            raise ImportError(
                "Install DOCX support with: pip install python-docx"
            ) from error

        uploaded_file.seek(0)
        word_document = WordDocument(uploaded_file)

        paragraphs = [
            paragraph.text.strip()
            for paragraph in word_document.paragraphs
            if paragraph.text.strip()
        ]

        uploaded_file.seek(0)

        seen_paragraphs: set[str] = set()

        text = self._remove_repeated_paragraphs(
            "\n\n".join(paragraphs),
            seen_paragraphs,
        )

        if not text:
            return [], 0

        return [
            Document(
                page_content=text,
                metadata={
                    "source": uploaded_file.name,
                    "page": 1,
                    "location": "Document",
                },
            )
        ], 1

    def _remove_exact_duplicate_chunks(
        self,
        chunks: list[Document],
    ) -> list[Document]:
        """
        Remove chunks that are effectively duplicates.

        Instead of requiring the entire normalized text to match,
        we use a prefix key so that nearly identical chunks with
        small trailing differences are collapsed.
        """
        unique_chunks: list[Document] = []
        seen: set[str] = set()
        prefix_len = 500  # tune as needed

        for chunk in chunks:
            normalized = self._normalize_for_comparison(
                chunk.page_content
            )

            if not normalized:
                continue

            key = normalized[:prefix_len]

            if key in seen:
                continue

            seen.add(key)
            unique_chunks.append(chunk)

        return unique_chunks

    def index_files(
        self,
        uploaded_files: list[Any],
    ) -> IndexResult:
        # Important: clear the current collection before indexing the
        # currently selected upload set. This prevents old chunks from
        # remaining in Chroma.
        self.vector_manager.clear()

        page_documents: list[Document] = []
        total_pages = 0
        total_files = 0
        errors: list[str] = []

        for uploaded_file in uploaded_files:
            try:
                lower_name = uploaded_file.name.lower()

                if lower_name.endswith(".pdf"):
                    documents, readable_pages = self._read_pdf(
                        uploaded_file
                    )
                elif lower_name.endswith(".docx"):
                    documents, readable_pages = self._read_docx(
                        uploaded_file
                    )
                else:
                    errors.append(
                        f"{uploaded_file.name}: unsupported file type."
                    )
                    continue

                if not documents:
                    errors.append(
                        f"{uploaded_file.name}: no readable text found."
                    )
                    continue

                page_documents.extend(documents)
                total_pages += readable_pages
                total_files += 1

            except Exception as error:
                errors.append(f"{uploaded_file.name}: {error}")

        chunks = self.text_splitter.split_documents(page_documents)

        chunks = self._remove_exact_duplicate_chunks(chunks)

        for chunk_index, chunk in enumerate(
            chunks,
            start=1,
        ):
            chunk.metadata["chunk_index"] = chunk_index
            chunk.metadata["index_version"] = INDEX_VERSION

        self.vector_manager.add_documents(chunks)

        return IndexResult(
            total_files=total_files,
            total_pages=total_pages,
            total_chunks=len(chunks),
            skipped_files=[],
            errors=errors,
        )

    def _deduplicate_retrieved_sources(
        self,
        sources: list[Document],
    ) -> list[Document]:
        unique: list[Document] = []
        seen: set[str] = set()

        for source in sources:
            normalized = self._normalize_for_comparison(
                source.page_content
            )

            if not normalized or normalized in seen:
                continue

            seen.add(normalized)
            unique.append(source)

        return unique

    def ask(
        self,
        question: str,
    ) -> AnswerResult:
        retrieved = self.vector_manager.similarity_search(
            question,
            k=8,
        )

        sources = self._deduplicate_retrieved_sources(
            retrieved
        )[:5]

        if not sources:
            return AnswerResult(
                answer=(
                    "I could not find relevant information "
                    "in the uploaded documents."
                ),
                sources=[],
            )

        context_sections: list[str] = []

        for number, source in enumerate(
            sources,
            start=1,
        ):
            source_name = source.metadata.get(
                "source",
                "Unknown document",
            )
            location = source.metadata.get(
                "location",
                "Unknown location",
            )

            context_sections.append(
                f"[Source {number}: {source_name}, {location}]\n"
                f"{source.page_content}"
            )

        context = "\n\n---\n\n".join(context_sections)

        prompt = f"""
You are RAGify, a document question-answering assistant.

Use only the document context below.

Requirements:
- Answer the question directly.
- Do not repeat the same fact.
- Do not invent information.
- When useful, mention the document and page.
- If the answer is absent, clearly say it was not found.

DOCUMENT CONTEXT:
{context}

QUESTION:
{question}

ANSWER:
""".strip()

        response = self.llm.invoke(prompt)
        answer = response.content

        if isinstance(answer, list):
            answer = " ".join(str(item) for item in answer)

        return AnswerResult(
            answer=str(answer).strip(),
            sources=sources,
        )
